import sys
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, fill):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_font(run, size=14, bold=False):
    run.font.name = 'TH SarabunPSK'
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts') if 'OxmlElement' in globals() else None
    if not rFonts:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'TH SarabunPSK')
    rFonts.set(qn('w:hAnsi'), 'TH SarabunPSK')
    rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    rFonts.set(qn('w:cs'), 'TH SarabunPSK')
    rPr.append(rFonts)
    run.font.size = Pt(size)
    run.bold = bold

def create_answer_sheet_4in1(output_path="กระดาษคำตอบ_4in1_Official.docx", school="โรงเรียนบ้านแม่ทราย(คุรุราษฎร์เจริญวิทย์)"):
    doc = Document()
    
    # 1. Strict Page Setup (Narrow margins to save space)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.4)
    section.bottom_margin = Cm(0.4)
    section.left_margin = Cm(0.4)
    section.right_margin = Cm(0.4)
    
    # 2. Master Table (2x2) - NO LOGO
    master = doc.add_table(rows=2, cols=2)
    master.style = 'Table Grid'
    master.autofit = False
    
    # Total Height = 29.7 - 0.8 = 28.9. Half is 14.45. Using 14.3 to be absolutely safe.
    for row in master.rows:
        row.height = Cm(14.3)

    for m_cell_el in master._element.xpath('.//w:tc'):
        from docx.table import _Cell
        cell_obj = _Cell(m_cell_el, master)
        cell_obj.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Reset default paragraph
        p_clear = cell_obj.paragraphs[0]
        p_clear.paragraph_format.line_spacing = 1.0
        p_clear.paragraph_format.space_after = Pt(0)
        
        # --- Header Area (No Logo, Centered) ---
        p1 = cell_obj.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(6)
        run1 = p1.add_run(school)
        set_font(run1, size=14, bold=True)
        
        p2 = cell_obj.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("วิชา................................................................................")
        set_font(run2, size=12)
        
        p3 = cell_obj.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run("ชื่อ............................................ เลขที่....... ชั้น........")
        set_font(run3, size=12)
        
        p_inst = cell_obj.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_inst.paragraph_format.space_after = Pt(2)
        run_inst = p_inst.add_run("ให้นักเรียนทำเครื่องหมาย X ลงในช่องที่เลือก")
        set_font(run_inst, size=10, bold=True)
        
        # --- Expanded Answer Grid (0.48cm height) ---
        grid = cell_obj.add_table(rows=23, cols=11)
        grid.style = 'Table Grid'
        grid.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r_row in grid.rows:
            r_row.height = Cm(0.48)

        pink_hex = "F4CCCC"
        headers = [
            ["ข้อ", "ก", "ข", "ค", "ง", "", "ข้อ", "ก", "ข", "ค", "ง"],
            ["", "A", "B", "C", "D", "", "", "A", "B", "C", "D"],
            ["", "1", "2", "3", "4", "", "", "1", "2", "3", "4"]
        ]
        
        for r in range(3):
            for c in range(11):
                cell = grid.cell(r, c)
                cell.text = headers[r][c]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if c != 5:
                    set_cell_shading(cell, pink_hex)
                for run in cell.paragraphs[0].runs:
                    set_font(run, size=9, bold=True)

        for i in range(1, 21):
            cl = grid.cell(i+2, 0); cl.text = str(i)
            cl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cl, pink_hex); set_font(cl.paragraphs[0].runs[0], size=9, bold=True)
            cr = grid.cell(i+2, 6); cr.text = str(i+20)
            cr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cr, pink_hex); set_font(cr.paragraphs[0].runs[0], size=9, bold=True)

        # Maximize width for symmetry
        ws = [0.8, 0.7, 0.7, 0.7, 0.7, 0.2, 0.8, 0.7, 0.7, 0.7, 0.7]
        for idx, w in enumerate(ws):
            for row in grid.rows:
                row.cells[idx].width = Cm(w)

    doc.save(output_path)
    print(f"Success: Created {output_path}")

if __name__ == "__main__":
    create_answer_sheet_4in1()
