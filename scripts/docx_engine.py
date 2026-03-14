import sys
import os
from docx import Document
from docx.shared import Twips, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class SiSomDocxEngine:
    def __init__(self, filename=None, font_size=16):
        if filename and os.path.exists(filename):
            self.doc = Document(filename)
        else:
            self.doc = Document()
            self._set_page_setup()
            self._set_tight_spacing()
            self.set_font('TH SarabunPSK', font_size)

    def _set_page_setup(self):
        """Standard A4 with precise margins (1418 Twips)"""
        section = self.doc.sections[0]
        section.page_height = Twips(16838)
        section.page_width = Twips(11906)
        section.top_margin = Twips(1418)
        section.bottom_margin = Twips(1418)
        section.left_margin = Twips(1418)
        section.right_margin = Twips(1418)

    def _set_tight_spacing(self):
        """Set default line spacing to Single (1.0) for tight layout."""
        style = self.doc.styles['Normal']
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

    def set_font(self, name='TH SarabunPSK', size=16):
        """Set default font with Thai support."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = name
        font.size = Pt(size)
        rPr = style._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:eastAsia'), name)
        rFonts.set(qn('w:cs'), name)
        rPr.append(rFonts)

    def add_official_header(self, subject, level, term, score, time, instruction, logo_path=None, school="โรงเรียนบ้านแม่ทราย(คุรุราษฎร์เจริญวิทย์)", area="สำนักงานเขตพื้นที่การศึกษาประถมศึกษาแพร่ เขต 1"):
        """Create the full official school header in a table with optional logo."""
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add Logo if provided
        if logo_path and os.path.exists(logo_path):
            run_logo = p.add_run()
            run_logo.add_picture(logo_path, width=Inches(0.8)) # Standard logo size
            p.add_run("\n") # Line break after logo

        header_text = (
            f"ข้อสอบวัดผลสัมฤทธิ์ปลายภาคเรียน  ภาคเรียนที่ {term}\n"
            f"วิชา{subject} ชั้น{level}  คะแนนเต็ม {score} คะแนน  เวลา {time} นาที\n"
            f"{school}  {area}\n"
            f"*************************************************************"
        )
        run = p.add_run(header_text)
        run.bold = True
        
        # ตารางชื่อ-สกุล (มาก่อนคำชี้แจง)
        table_info = self.doc.add_table(rows=1, cols=2)
        table_info.style = 'Table Grid'
        table_info.cell(0, 0).text = "ชื่อ–สกุล  ........................................................."
        table_info.cell(0, 1).text = "ชั้น  ........  เลขที่  ........"
        
        # คำชี้แจง (มาทีหลังชื่อ-สกุล)
        p_inst = self.doc.add_paragraph()
        p_inst.paragraph_format.space_before = Pt(6)
        run_inst = p_inst.add_run(f"คำชี้แจง  {instruction}")
        run_inst.bold = True
        self.doc.add_paragraph("") # Spacer

    def add_answer_key_table(self, answers):
        """Create a grid table for answer keys (5 columns)."""
        self.doc.add_section(WD_SECTION.NEW_PAGE)
        self.doc.add_paragraph("-" * 80)
        p = self.doc.add_paragraph("ส่วนที่ 2  เฉลยข้อสอบ (Answer Key)")
        p.runs[0].bold = True
        
        num_ans = len(answers)
        rows = (num_ans + 4) // 5
        table = self.doc.add_table(rows=rows, cols=5)
        table.style = 'Table Grid'
        
        for i, ans in enumerate(answers):
            row = i // 5
            col = i % 5
            table.cell(row, col).text = f"{i+1}. {ans}"

    def set_2_columns_continuous(self):
        """Add a continuous section break and set to 2 columns."""
        new_section = self.doc.add_section(WD_SECTION.CONTINUOUS)
        sectPr = new_section._sectPr
        cols = sectPr.xpath('./w:cols')
        if not cols:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        else:
            cols = cols[0]
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720') # 0.5 inch gap
        return new_section

    def add_header_center(self, text, bold=False):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold and p.runs:
            p.runs[0].bold = True
        return p

    def save(self, filename):
        self.doc.save(filename)
        print(f"✅ Document saved: {filename}")

if __name__ == "__main__":
    print("Si-Som DOCX Engine v2.0 (Exam Ready)")
